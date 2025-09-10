from io import BytesIO
from typing import Any

from docx import Document
from PyPDF2 import PdfReader


def extract_pdf_meta(data: bytes) -> dict[str, Any]:
    stream = BytesIO(data)
    reader = PdfReader(stream)
    info = reader.metadata or {}
    pages = len(reader.pages)
    author = str(info.get("/Author")) if info else None
    title = str(info.get("/Title")) if info else None
    producer = str(info.get("/Producer")) if info else None
    created = str(info.get("/CreationDate")) if info else None
    return {
        "pages": pages,
        "author": author,
        "title": title,
        "producer": producer,
        "created": created,
    }


def extract_docx_meta(data: bytes) -> dict[str, Any]:
    stream = BytesIO(data)
    doc = Document(stream)
    core = doc.core_properties
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "title": core.title,
        "author": core.author,
        "created": str(core.created),
    }
